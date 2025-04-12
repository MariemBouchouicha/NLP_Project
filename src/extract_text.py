import os
import pandas as pd
import pdfplumber
from docx import Document

def extract_text(file_path):
    """
    Extract text from .docx, .pdf, .csv, .xlsx, .xls, .xlsm files.
    Returns everything as a plain string with '|' as table separator.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[-1].lower()

    if ext == ".docx":
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        tables = "\n".join(
            "\n".join(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            for table in doc.tables
        )
        return f"{text}\n\n--- Table ---\n{tables}".strip()

    if ext in [".xlsx", ".xls", ".xlsm"]:
        sheets = pd.read_excel(file_path, sheet_name=None)
        return "\n".join(
            f"--- Sheet: {name} ---\n{df.to_csv(sep='|', index=False)}"
            for name, df in sheets.items()
        )

    if ext == ".csv":
        return pd.read_csv(file_path).to_csv(sep="|", index=False)

    if ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    raise ValueError(f"Unsupported file type: {ext}")
